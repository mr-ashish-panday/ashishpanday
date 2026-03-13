from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


TITLE = "3D Portfolio Website Project Report"
AUTHOR = "Ashish Panday"
PROJECT_TYPE = "Full-Stack Web Development Project"
GENERATED_ON = "March 13, 2026"
SUBTITLE = "Frontend, Backend, API Integration, and Deployment Support"
LIVE_LINK = "https://ashishpanday.vercel.app/"
ABSTRACT = (
    "This report presents the design and development of a full-stack 3D portfolio "
    "website that combines immersive frontend presentation with practical backend "
    "features such as API routing, validation, contact form handling, storage, and "
    "deployment support."
)

INTRO = [
    "This project is a modern full-stack personal portfolio website developed to present academic, technical, and research work in a visually engaging format. The application combines a 3D animated frontend with a functional backend, turning a static portfolio into a complete web application.",
    "The project was designed to demonstrate both creative frontend development and practical backend implementation. It shows not only user interface and animation skills, but also server-side programming, form processing, data validation, API integration, and deployment readiness.",
]

OBJECTIVES = [
    "To create a visually impressive 3D portfolio website.",
    "To implement responsive frontend sections for content presentation.",
    "To add backend functionality for real user interaction.",
    "To create API routes for portfolio data and message handling.",
    "To validate contact form input before storing data.",
    "To prepare the project for deployment on Vercel.",
    "To make the project suitable as a complete college-level full-stack application.",
]

TECH_STACK = [
    (
        "Frontend Technologies",
        [
            "React",
            "Vite",
            "Tailwind CSS",
            "GSAP",
            "Three.js",
            "React Three Fiber",
            "Drei",
            "Lenis",
            "Motion",
        ],
    ),
    (
        "Backend Technologies",
        [
            "Node.js",
            "Express.js",
            "CORS",
            "REST API architecture",
        ],
    ),
    (
        "Deployment and Storage",
        [
            "Vercel",
            "Vercel Functions",
            "Vercel Blob Storage",
        ],
    ),
]

SYSTEM_OVERVIEW = [
    "The system is divided into two major parts: the frontend and the backend. The frontend is responsible for the visual presentation of the portfolio, while the backend handles data communication, API responses, validation, and message storage.",
    "The main frontend sections are Hero, About, Services, Works, Contact Summary, and Contact. The backend exposes routes such as /api/portfolio, /api/health, /api/messages, and /api/contact. The frontend communicates with these routes through HTTP requests.",
    "The project also supports two presentation modes: a dynamic 3D React-based experience and a lightweight static HTML version. A visible toggle allows users to switch between the dynamic and static views, making the website flexible for both immersive presentation and minimal reading mode.",
]

FRONTEND = [
    "The frontend was built using React and Vite, with Tailwind CSS for styling and GSAP for animations. Three.js through React Three Fiber and Drei was used to create the 3D visual experience in the Hero section.",
    "A smooth loading experience was implemented using useProgress from Drei, and scroll behavior was enhanced with Lenis. Different sections of the portfolio present the user's research profile, services, projects, and contact information in a responsive layout.",
    "In addition to the animated 3D version, the project includes a static minimal page built with HTML and CSS. The main dynamic website provides a Minimal toggle, and the static page provides a 3D View toggle, so users can move back and forth between the two experiences.",
]

FRONTEND_FEATURES = [
    "Animated hero section with 3D graphics.",
    "Smooth transitions and scroll-based motion.",
    "Responsive layout across desktop and mobile devices.",
    "Dynamic services and works sections.",
    "Dual presentation modes with toggle support between static and dynamic views.",
    "Interactive contact section with backend integration.",
    "Frontend fallback data when the backend is unavailable.",
]

BACKEND = [
    "The backend was implemented using Node.js and Express.js for local development. Its purpose is to process requests from the frontend, manage data exchange, validate user input, and support deployment-ready API behavior.",
    "The backend logic was structured into service and storage layers. Shared backend modules were created so the same logic can be used for both local Express routes and deployed Vercel serverless functions.",
]

API_ENDPOINTS = [
    "GET /api/portfolio: provides profile information, services, and project data.",
    "GET /api/health: returns backend status, submission count, and storage mode.",
    "GET /api/messages: returns stored contact form submissions.",
    "POST /api/contact: receives, validates, and stores contact form messages.",
]

CONTACT_SYSTEM = [
    "The Contact section includes a real form where users submit their name, email, subject, and message. When the form is submitted, the frontend sends the data to the backend using an API request.",
    "The backend validates each field before saving the message. If any field is invalid, the backend returns field-specific error messages to the frontend. If all fields are valid, the message is stored and a success response is returned to the user.",
]

VALIDATION_RULES = [
    "Name must contain meaningful text.",
    "Email must be in valid email format.",
    "Subject must not be too short.",
    "Message must contain sufficient content.",
]

STORAGE = [
    "During local development, submitted messages are stored in a JSON file. This makes testing simple and allows inspection of saved data during development.",
    "For deployment, the project was upgraded to support Vercel Blob Storage because local file writing is not reliable in serverless hosting environments. This allows the deployed contact form to keep working correctly after publication.",
]

DEPLOYMENT = [
    "The project was prepared for deployment on Vercel. Dedicated API handlers were created so that the same backend features can run as Vercel Functions in production.",
    "This means the website can be shared through a live link, and the professor or any other viewer can use the contact form successfully, provided the Vercel project is connected to Blob storage.",
]

WORKFLOW = [
    "The user opens the website.",
    "The frontend loads the portfolio sections and requests backend data.",
    "The backend returns portfolio data and health information.",
    "The user fills in the contact form and submits a message.",
    "The backend validates the data and stores the message.",
    "The frontend displays a success or error response.",
]

ACHIEVEMENTS = [
    "Built a visually rich 3D portfolio website.",
    "Added a working backend with multiple API endpoints.",
    "Implemented a real contact form system.",
    "Added input validation and error handling.",
    "Integrated frontend and backend successfully.",
    "Added deployment-ready support for Vercel.",
    "Supported both local file storage and Vercel Blob storage.",
]

CHALLENGES = [
    "The original project was frontend-only, so backend integration had to be added without disturbing the visual structure.",
    "The contact form initially failed when only the frontend was running and the backend server was not active.",
    "Local JSON storage was useful for development but unsuitable for permanent cloud deployment.",
    "The backend needed to be redesigned to work correctly on Vercel using serverless API functions and persistent storage.",
]

CONCLUSION = [
    "This project successfully combines modern frontend design with practical backend engineering. The frontend offers a strong visual experience through animation, responsiveness, and 3D graphics, while the backend adds real-world functionality through APIs, validation, storage, and deployment support.",
    "As a result, the project is not only a creative portfolio but also a meaningful full-stack web application suitable for academic presentation and real deployment.",
]

FUTURE_IMPROVEMENTS = [
    "Admin dashboard to view submitted messages.",
    "Authentication system for secure access.",
    "Database integration such as MongoDB or PostgreSQL.",
    "Automatic email notification after form submission.",
    "Analytics dashboard for traffic and message insights.",
    "Content management panel for updating projects dynamically.",
]

REFERENCES = [
    "React Documentation",
    "Vite Documentation",
    "Tailwind CSS Documentation",
    "GSAP Documentation",
    "Three.js Documentation",
    "React Three Fiber Documentation",
    "Node.js Documentation",
    "Express.js Documentation",
    "Vercel Documentation",
    "Vercel Blob Documentation",
]


def add_title_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(20)

    document.add_paragraph()
    document.add_paragraph()

    for label, value in [
        ("Submitted By", AUTHOR),
        ("Project Type", PROJECT_TYPE),
        ("Generated On", GENERATED_ON),
        ("Live Website", LIVE_LINK),
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{label}: {value}")
        run.font.size = Pt(12)

    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraphs(document: Document, paragraphs) -> None:
    for text in paragraphs:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(document: Document, items) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def build_docx(output_path: Path) -> None:
    document = Document()
    add_title_page(document)

    add_heading(document, "1. Introduction")
    add_paragraphs(document, INTRO)

    add_heading(document, "2. Objectives")
    add_bullets(document, OBJECTIVES)

    add_heading(document, "3. Technologies Used")
    for heading, items in TECH_STACK:
        add_heading(document, heading, level=2)
        add_bullets(document, items)

    add_heading(document, "4. System Overview")
    add_paragraphs(document, SYSTEM_OVERVIEW)

    add_heading(document, "5. Frontend Implementation")
    add_paragraphs(document, FRONTEND)
    add_heading(document, "Main Frontend Features", level=2)
    add_bullets(document, FRONTEND_FEATURES)

    add_heading(document, "6. Backend Implementation")
    add_paragraphs(document, BACKEND)
    add_heading(document, "Main API Endpoints", level=2)
    add_bullets(document, API_ENDPOINTS)

    add_heading(document, "7. Contact Form Functionality")
    add_paragraphs(document, CONTACT_SYSTEM)
    add_heading(document, "Validation Rules", level=2)
    add_bullets(document, VALIDATION_RULES)

    add_heading(document, "8. Data Storage")
    add_paragraphs(document, STORAGE)

    add_heading(document, "9. Deployment Support")
    add_paragraphs(document, DEPLOYMENT)

    add_heading(document, "10. Project Workflow")
    add_bullets(document, WORKFLOW)

    add_heading(document, "11. Key Achievements")
    add_bullets(document, ACHIEVEMENTS)

    add_heading(document, "12. Challenges Faced")
    add_bullets(document, CHALLENGES)

    add_heading(document, "13. Conclusion")
    add_paragraphs(document, CONCLUSION)

    add_heading(document, "14. Future Improvements")
    add_bullets(document, FUTURE_IMPROVEMENTS)

    add_heading(document, "15. References")
    add_bullets(document, REFERENCES)

    document.save(output_path)


def build_pdf(output_path: Path) -> None:
    primary = HexColor("#0f172a")
    accent = HexColor("#2563eb")
    muted = HexColor("#64748b")
    border = HexColor("#cbd5e1")
    soft_bg = HexColor("#f8fafc")
    soft_accent = HexColor("#eff6ff")

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=30,
            alignment=TA_CENTER,
            textColor=primary,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportKicker",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
            textColor=accent,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportSubtitle",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=11.5,
            leading=16,
            alignment=TA_CENTER,
            textColor=muted,
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportHeading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=20,
            textColor=primary,
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportSubHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=15,
            textColor=accent,
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportBody",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10.5,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            textColor=primary,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportMeta",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            alignment=TA_CENTER,
            textColor=primary,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportBullet",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10.5,
            leading=15,
            alignment=TA_LEFT,
            leftIndent=14,
            textColor=primary,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportSummary",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=10.5,
            leading=16,
            alignment=TA_JUSTIFY,
            textColor=primary,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportLabel",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            alignment=TA_LEFT,
            textColor=muted,
            uppercase=True,
        )
    )

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )
    story = []

    meta_table = Table(
        [
            ["SUBMITTED BY", AUTHOR],
            ["PROJECT TYPE", PROJECT_TYPE],
            ["GENERATED ON", GENERATED_ON],
            ["LIVE WEBSITE", LIVE_LINK],
        ],
        colWidths=[1.7 * inch, 4.8 * inch],
    )
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), soft_bg),
                ("BOX", (0, 0), (-1, -1), 0.8, border),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, border),
                ("TEXTCOLOR", (0, 0), (0, -1), muted),
                ("TEXTCOLOR", (1, 0), (1, -1), primary),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 10.5),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    summary_box = Table(
        [[Paragraph(ABSTRACT, styles["ReportSummary"])]],
        colWidths=[6.5 * inch],
    )
    summary_box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), soft_accent),
                ("BOX", (0, 0), (-1, -1), 0.8, accent),
                ("TOPPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ("LEFTPADDING", (0, 0), (-1, -1), 14),
                ("RIGHTPADDING", (0, 0), (-1, -1), 14),
            ]
        )
    )

    story.append(Spacer(1, 1.4 * inch))
    story.append(Paragraph("PROJECT REPORT", styles["ReportKicker"]))
    story.append(Paragraph(TITLE, styles["ReportTitle"]))
    story.append(Paragraph(SUBTITLE, styles["ReportSubtitle"]))
    story.append(
        HRFlowable(width="28%", thickness=1.4, color=accent, spaceBefore=2, spaceAfter=22)
    )
    story.append(meta_table)
    story.append(Spacer(1, 0.45 * inch))
    story.append(summary_box)
    story.append(PageBreak())

    def section_heading(title):
        return KeepTogether(
            [
                Paragraph(title, styles["ReportHeading"]),
                HRFlowable(
                    width="100%",
                    thickness=0.8,
                    color=border,
                    spaceBefore=0,
                    spaceAfter=10,
                ),
            ]
        )

    def add_section(title, paragraphs=None, bullets=None):
        story.append(section_heading(title))
        if paragraphs:
            for text in paragraphs:
                story.append(Paragraph(text, styles["ReportBody"]))
        if bullets:
            flow_items = [
                ListItem(Paragraph(item, styles["ReportBullet"])) for item in bullets
            ]
            story.append(
                ListFlowable(
                    flow_items,
                    bulletType="bullet",
                    start="circle",
                    leftIndent=12,
                    bulletFontName="Helvetica",
                    bulletFontSize=10,
                    bulletColor=accent,
                )
            )
            story.append(Spacer(1, 8))

    abstract_box = Table(
        [[Paragraph(ABSTRACT, styles["ReportSummary"])]],
        colWidths=[6.5 * inch],
    )
    abstract_box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), soft_bg),
                ("BOX", (0, 0), (-1, -1), 0.8, border),
                ("TOPPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ("LEFTPADDING", (0, 0), (-1, -1), 14),
                ("RIGHTPADDING", (0, 0), (-1, -1), 14),
            ]
        )
    )
    story.append(Paragraph("ABSTRACT", styles["ReportSubHeading"]))
    story.append(abstract_box)
    story.append(Spacer(1, 14))

    add_section("1. Introduction", paragraphs=INTRO)
    add_section("2. Objectives", bullets=OBJECTIVES)

    story.append(section_heading("3. Technologies Used"))
    for heading, items in TECH_STACK:
        story.append(Paragraph(heading, styles["ReportSubHeading"]))
        flow_items = [ListItem(Paragraph(item, styles["ReportBullet"])) for item in items]
        story.append(
            ListFlowable(
                flow_items,
                bulletType="bullet",
                start="circle",
                leftIndent=12,
                bulletFontName="Helvetica",
                bulletFontSize=10,
                bulletColor=accent,
            )
        )
        story.append(Spacer(1, 8))

    add_section("4. System Overview", paragraphs=SYSTEM_OVERVIEW)
    add_section("5. Frontend Implementation", paragraphs=FRONTEND)
    story.append(Paragraph("Main Frontend Features", styles["ReportSubHeading"]))
    story.append(
        ListFlowable(
            [ListItem(Paragraph(item, styles["ReportBullet"])) for item in FRONTEND_FEATURES],
            bulletType="bullet",
            start="circle",
            leftIndent=12,
            bulletFontName="Helvetica",
            bulletFontSize=10,
            bulletColor=accent,
        )
    )
    story.append(Spacer(1, 8))

    add_section("6. Backend Implementation", paragraphs=BACKEND)
    story.append(Paragraph("Main API Endpoints", styles["ReportSubHeading"]))
    story.append(
        ListFlowable(
            [ListItem(Paragraph(item, styles["ReportBullet"])) for item in API_ENDPOINTS],
            bulletType="bullet",
            start="circle",
            leftIndent=12,
            bulletFontName="Helvetica",
            bulletFontSize=10,
            bulletColor=accent,
        )
    )
    story.append(Spacer(1, 8))

    add_section("7. Contact Form Functionality", paragraphs=CONTACT_SYSTEM)
    story.append(Paragraph("Validation Rules", styles["ReportSubHeading"]))
    story.append(
        ListFlowable(
            [ListItem(Paragraph(item, styles["ReportBullet"])) for item in VALIDATION_RULES],
            bulletType="bullet",
            start="circle",
            leftIndent=12,
            bulletFontName="Helvetica",
            bulletFontSize=10,
            bulletColor=accent,
        )
    )
    story.append(Spacer(1, 8))

    add_section("8. Data Storage", paragraphs=STORAGE)
    add_section("9. Deployment Support", paragraphs=DEPLOYMENT)
    add_section("10. Project Workflow", bullets=WORKFLOW)
    add_section("11. Key Achievements", bullets=ACHIEVEMENTS)
    add_section("12. Challenges Faced", bullets=CHALLENGES)
    add_section("13. Conclusion", paragraphs=CONCLUSION)
    add_section("14. Future Improvements", bullets=FUTURE_IMPROVEMENTS)
    add_section("15. References", bullets=REFERENCES)

    def draw_first_page(canvas, _doc):
        canvas.saveState()
        canvas.setStrokeColor(border)
        canvas.setLineWidth(1)
        canvas.line(0.85 * inch, 0.7 * inch, A4[0] - 0.85 * inch, 0.7 * inch)
        canvas.setFillColor(muted)
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(A4[0] / 2, 0.48 * inch, AUTHOR)
        canvas.restoreState()

    def draw_later_pages(canvas, _doc):
        canvas.saveState()
        canvas.setStrokeColor(border)
        canvas.setLineWidth(0.7)
        canvas.line(0.85 * inch, A4[1] - 0.6 * inch, A4[0] - 0.85 * inch, A4[1] - 0.6 * inch)
        canvas.line(0.85 * inch, 0.6 * inch, A4[0] - 0.85 * inch, 0.6 * inch)
        canvas.setFillColor(muted)
        canvas.setFont("Helvetica", 9)
        canvas.drawString(0.85 * inch, 0.38 * inch, TITLE)
        canvas.drawRightString(
            A4[0] - 0.85 * inch, 0.38 * inch, f"Page {canvas.getPageNumber() - 1}"
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_first_page, onLaterPages=draw_later_pages)


def main() -> None:
    output_dir = Path(__file__).resolve().parent
    docx_path = output_dir / "3D_Portfolio_Project_Report.docx"
    pdf_path = output_dir / "3D_Portfolio_Project_Report.pdf"

    build_docx(docx_path)
    build_pdf(pdf_path)

    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
